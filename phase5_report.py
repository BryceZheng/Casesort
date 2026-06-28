#!/usr/bin/env python3
"""
Casesort Phase 5 - 法律分析报告 Word 文档生成工具

用法：
  python3 phase5_report.py --input report.json --output 案例分析报告.docx

JSON 结构：
{
  "title": "报告标题",
  "subtitle": "副标题（如数据来源与时间范围）",
  "sections": [
    {
      "heading": "一、数据概览",
      "level": 1,
      "content": [
        {"type": "paragraph", "text": "正文段落"},
        {"type": "bullet",    "items": ["要点1", "要点2"]},
        {"type": "table",     "headers": ["列1","列2"], "rows": [["值","值"]]}
      ]
    }
  ]
}
"""

import argparse
import json
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("请先安装 python-docx：pip3 install python-docx", file=sys.stderr)
    sys.exit(1)


# ── 样式常量 ────────────────────────────────────────────────
THEME_COLOR   = RGBColor(0x2F, 0x54, 0x96)   # 深蓝色标题
TABLE_HEADER  = RGBColor(0x2F, 0x54, 0x96)
FONT_CHINESE  = "宋体"
FONT_LATIN    = "Times New Roman"


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_paragraph(doc, text: str, level: int = 0, bold: bool = False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if level > 0:
        p.paragraph_format.left_indent = Cm(0.5 * level)
    run = p.add_run(text)
    run.font.name = FONT_CHINESE
    run.font.size = Pt(10.5)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CHINESE)
    return p


def add_heading(doc, text: str, level: int):
    p = doc.add_heading("", level=level)
    p.clear()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT_CHINESE
    run.font.color.rgb = THEME_COLOR
    run.font.size = Pt(14 if level == 1 else 12)
    run.bold = True
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CHINESE)
    return p


def add_bullet(doc, items: list):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(item)
        run.font.name = FONT_CHINESE
        run.font.size = Pt(10.5)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CHINESE)


def add_table(doc, headers: list, rows: list):
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"

    # 表头行
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, "2F5496")
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = FONT_CHINESE
        run.font.size = Pt(10)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CHINESE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "EBF3FB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = FONT_CHINESE
            run.font.size = Pt(10)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CHINESE)

    doc.add_paragraph()


def build_doc(data: dict, output_path: str):
    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # 封面标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(20)
    title_p.paragraph_format.space_after  = Pt(6)
    run = title_p.add_run(data.get("title", "法律案例分析报告"))
    run.font.name = FONT_CHINESE
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = THEME_COLOR
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CHINESE)

    if data.get("subtitle"):
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.paragraph_format.space_after = Pt(16)
        run2 = sub_p.add_run(data["subtitle"])
        run2.font.name = FONT_CHINESE
        run2.font.size = Pt(11)
        run2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        run2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CHINESE)

    doc.add_paragraph()

    # 正文各章节
    for section in data.get("sections", []):
        heading_text = section.get("heading", "")
        level = section.get("level", 1)
        add_heading(doc, heading_text, level)

        for block in section.get("content", []):
            btype = block.get("type", "paragraph")
            if btype == "paragraph":
                add_paragraph(doc, block.get("text", ""))
            elif btype == "bullet":
                add_bullet(doc, block.get("items", []))
            elif btype == "table":
                add_table(doc, block.get("headers", []), block.get("rows", []))

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"✅ 报告已生成：{output_path}")


def main():
    parser = argparse.ArgumentParser(description="Casesort Phase 5 报告生成")
    parser.add_argument("--input",  required=True, help="报告 JSON 文件路径")
    parser.add_argument("--output", required=True, help="输出 Word 文件路径（.docx）")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"错误：输入文件不存在：{args.input}", file=sys.stderr)
        sys.exit(1)

    with open(args.input, "r", encoding="utf-8") as f:
        data = json.load(f)

    build_doc(data, args.output)


if __name__ == "__main__":
    main()
